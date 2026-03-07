import os
import re
import requests
from typing import Optional
from pathlib import Path

def extract_from_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n\n".join(text)
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}")

def extract_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}")

def extract_from_txt(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Text extraction failed: {e}")

def extract_from_image(file_path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        if not text.strip():
            raise ValueError("No text found in image (OCR returned empty result)")
        return text
    except ImportError:
        raise ValueError("OCR libraries not installed. Install pytesseract and Pillow.")
    except Exception as e:
        raise ValueError(f"Image OCR failed: {e}")

def extract_from_url(url: str) -> str:
    try:
        from bs4 import BeautifulSoup
        headers = {
            'User-Agent': 'Mozilla/5.0 (compatible; QuizBot/1.0)'
        }
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, 'html.parser')
        
        # Remove scripts, styles, nav, footer
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'advertisement']):
            tag.decompose()
        
        # Try to find main content
        main = soup.find('main') or soup.find('article') or soup.find('div', {'id': 'content'}) or soup.body
        
        if main:
            text = main.get_text(separator='\n', strip=True)
        else:
            text = soup.get_text(separator='\n', strip=True)
        
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = '\n'.join(lines)
        
        if len(text) < 100:
            raise ValueError("Extracted content too short. The page may be dynamically rendered.")
        
        return text
    except requests.RequestException as e:
        raise ValueError(f"Failed to fetch URL: {e}")
    except Exception as e:
        raise ValueError(f"URL extraction failed: {e}")

def extract_content(file_path: Optional[str] = None, file_type: Optional[str] = None, url: Optional[str] = None) -> str:
    if url:
        content = extract_from_url(url)
    elif file_path and file_type:
        ext = file_type.lower()
        if ext in ['pdf']:
            content = extract_from_pdf(file_path)
        elif ext in ['docx']:
            content = extract_from_docx(file_path)
        elif ext in ['txt', 'text']:
            content = extract_from_txt(file_path)
        elif ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
            content = extract_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    else:
        raise ValueError("Either file_path+file_type or url must be provided")
    
    return clean_text(content)

def clean_text(text: str) -> str:
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    # Remove control characters except newlines
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text.strip()
